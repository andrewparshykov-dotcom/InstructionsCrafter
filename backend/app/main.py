"""FastAPI application for the InstructionsCrafter backend."""

import io
import os
import re
import secrets
from datetime import date

from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse
from starlette.exceptions import HTTPException as StarletteHTTPException

load_dotenv()

SHARED_PASSWORD = os.getenv("SHARED_PASSWORD", "")
MAX_VIDEO_SIZE_MB = int(os.getenv("MAX_VIDEO_SIZE_MB", "500"))
MAX_VIDEO_SIZE_BYTES = MAX_VIDEO_SIZE_MB * 1024 * 1024

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

app = FastAPI(title="InstructionsCrafter API")


@app.exception_handler(StarletteHTTPException)
async def http_exception_handler(request: Request, exc: StarletteHTTPException):
    # Architecture spec wants {"error": "..."} instead of FastAPI's default {"detail": "..."}.
    return JSONResponse(
        status_code=exc.status_code,
        content={"error": exc.detail},
    )


@app.get("/api/health")
async def health():
    return {"status": "ok"}


@app.post("/api/generate")
async def generate(
    video: UploadFile = File(...),
    title: str = Form(...),
    password: str = Form(...),
):
    if not SHARED_PASSWORD or not secrets.compare_digest(password, SHARED_PASSWORD):
        raise HTTPException(status_code=401, detail="Unauthorized")

    if not (1 <= len(title) <= 200):
        raise HTTPException(status_code=400, detail="Title must be 1 to 200 characters")

    if video.size is not None and video.size > MAX_VIDEO_SIZE_BYTES:
        raise HTTPException(status_code=413, detail="Video exceeds maximum size")
    if video.size == 0:
        raise HTTPException(status_code=400, detail="Empty video file")

    # Phase 1 placeholder — real pipeline arrives in later phases.
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated on {date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph(
        "This is a placeholder document produced by Phase 1 of the backend. "
        "Later phases replace this with real transcribed steps and screenshots."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{_sanitize_title(title)}_{date.today().isoformat()}.docx"
    return StreamingResponse(
        buffer,
        media_type=DOCX_MEDIA_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _sanitize_title(title: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9 \-]", "_", title).strip()
    return cleaned or "document"
