"""Generate the InstructionsCrafter Word template.

Writes a python-docx-template (docxtpl)-compatible .docx file at
backend/templates/instruction_template.docx.

The template provides these placeholders:
    {{ title }}              -- document title (top of page)
    {{ date }}               -- formatted generation date
    {%p for step in steps %} -- begin per-step block
        {{ loop.index }}     -- current step number (1, 2, 3, ...)
        {{ step.instruction }}
        {{ step.image }}     -- replaced with an InlineImage by docxtpl
        {{ step.caption }}
    {%p endfor %}            -- end per-step block

The team can later open the resulting .docx in Microsoft Word and
adjust fonts, colors, or spacing -- but the Jinja tags ({{ ... }} and
{%p ... %}) must be preserved or rendering will break.

Re-running this script overwrites any manual edits to the .docx, so
adjust the script (or stop using it) once the template is finalized.

Run from the project root:
    python backend/scripts/build_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OUTPUT = (
    Path(__file__).resolve().parent.parent
    / "templates"
    / "instruction_template.docx"
)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Page setup: A4 landscape with 0.635 cm (0.25 in) uniform margins.
    # Landscape orientation gives screenshots more horizontal room; the very
    # tight margins maximize content area for screen-only viewing (this
    # document is digital, not printed).
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(0.635)
    section.bottom_margin = Cm(0.635)
    section.left_margin = Cm(0.635)
    section.right_margin = Cm(0.635)

    # Title -- 24pt bold, centered (Title style).
    title_p = doc.add_paragraph(style=doc.styles["Title"])
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run("{{ title }}")

    # Subtitle -- italic, centered (Subtitle style already styles this).
    subtitle_p = doc.add_paragraph(style=doc.styles["Subtitle"])
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.add_run("Generated on {{ date }}")

    # Optional Overview block -- wrapped in {%p if introduction %} so it
    # only renders when the AI-generated introduction is non-empty. On
    # rule-based-fallback runs the introduction is "" and the entire
    # block (heading + paragraph) disappears.
    doc.add_paragraph("{%p if introduction %}")

    overview_heading_p = doc.add_paragraph(style=doc.styles["Heading 2"])
    overview_heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    overview_heading_p.paragraph_format.keep_with_next = True
    overview_heading_p.add_run("Overview")

    intro_p = doc.add_paragraph(style=doc.styles["Body Text"])
    intro_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_p.paragraph_format.space_after = Pt(12)
    intro_p.add_run("{{ introduction }}")

    doc.add_paragraph("{%p endif %}")

    # For-loop opener -- {%p ... %} tells docxtpl to delete the entire
    # paragraph after processing, so no blank line appears in output.
    doc.add_paragraph("{%p for step in steps %}")

    # Step heading -- "Step N" (Heading 2 style).
    # `keep_with_next` glues the heading to the instruction so they cannot
    # be split across pages.
    heading_p = doc.add_paragraph(style=doc.styles["Heading 2"])
    heading_p.paragraph_format.keep_with_next = True
    heading_p.add_run("Step {{ loop.index }}")

    # Instruction body -- justified, 12pt space-after (Body Text style).
    # `keep_with_next` glues the instruction to the image below.
    instruction_p = doc.add_paragraph(style=doc.styles["Body Text"])
    instruction_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    instruction_p.paragraph_format.space_after = Pt(12)
    instruction_p.paragraph_format.keep_with_next = True
    instruction_p.add_run("{{ step.instruction }}")

    # Image placeholder -- centered. docxtpl replaces {{ step.image }} with
    # the InlineImage object passed in the rendering context.
    # `keep_with_next` glues the image to its caption.
    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.keep_with_next = True
    image_p.add_run("{{ step.image }}")

    # Caption -- italic, centered (Caption style is italic by default).
    caption_p = doc.add_paragraph(style=doc.styles["Caption"])
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.add_run("{{ step.caption }}")

    # Blank line for visual breathing room between steps.
    doc.add_paragraph()

    # For-loop closer.
    doc.add_paragraph("{%p endfor %}")

    doc.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")
    print(f"Size:  {OUTPUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
